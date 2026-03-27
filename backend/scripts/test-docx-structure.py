#!/usr/bin/env python3
"""
测试 Word 文档结构
"""

from docx import Document

doc_path = '/Users/rabbit-y/Documents/消防装备供应商资料/供应商/消防设备供应商目录（2025）/消防设备供应商目录word版（清晰）/消防设备供应商目录（2025）_1-100.docx'

doc = Document(doc_path)

print(f'📊 文档统计:\n')
print(f'   总段落数: {len(doc.paragraphs)}')
print(f'   总表格数: {len(doc.tables)}')

print(f'\n📝 前30段内容:\n')
for i, para in enumerate(doc.paragraphs[:30]):
    text = para.text.strip()
    if text:
        print(f'{i+1}. {text[:150]}')

if doc.tables:
    print(f'\n📋 第一个表格结构:\n')
    table = doc.tables[0]
    print(f'   行数: {len(table.rows)}')
    print(f'   列数: {len(table.columns) if table.rows else 0}')
    
    print(f'\n   前 5 行内容:\n')
    for i, row in enumerate(table.rows[:5]):
        cells = [cell.text.strip() for cell in row.cells]
        print(f'   行 {i+1}: {cells}')
