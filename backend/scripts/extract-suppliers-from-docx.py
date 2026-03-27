#!/usr/bin/env python3
"""
从 Word 文档提取展会供应商数据
使用本地 Ollama 模型（qwen2.5:14b）解析
"""

from docx import Document
import json
import re
import ollama

def extract_suppliers_from_docx(doc_path):
    """提取供应商信息"""
    doc = Document(doc_path)
    suppliers = []
    
    current_supplier = {}
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            # 空行可能表示供应商条目结束
            if current_supplier:
                suppliers.append(current_supplier)
                current_supplier = {}
            continue
        
        # 检测字段（通过关键词）
        if '公司名称' in text or 'Company Name' in text:
            # 提取公司名称
            name = text.split(':', 1)[-1].strip() if ':' in text else text
            current_supplier['company_name'] = name
        
        elif '地址' in text or 'Address' in text:
            address = text.split(':', 1)[-1].strip() if ':' in text else text
            current_supplier['address'] = address
        
        elif '邮编' in text or 'P.C' in text or 'Postcode' in text:
            postcode = re.sub(r'[^\d]', '', text)
            current_supplier['postcode'] = postcode
        
        elif '电话' in text or 'Tel' in text or 'Phone' in text:
            phone = text.split(':', 1)[-1].strip() if ':' in text else text
            phone = re.sub(r'[^0-9\-\s\+\(\)]', '', phone).strip()
            current_supplier['phone'] = phone
        
        elif '传真' in text or 'Fax' in text:
            fax = text.split(':', 1)[-1].strip() if ':' in text else text
            current_supplier['fax'] = fax
        
        elif '电邮' in text or 'E-mail' in text or 'Email' in text:
            email = text.split(':', 1)[-1].strip() if ':' in text else text
            email = re.sub(r'[^\w@\.\-]', '', email)
            current_supplier['email'] = email
        
        elif '网址' in text or 'URL' in text or 'Website' in text or 'http' in text:
            url = text.split(':', 1)[-1].strip() if ':' in text else text
            url = url.replace(' ', '').strip()
            if url.startswith('http'):
                current_supplier['website'] = url
        
        elif '联系人' in text or 'Contact' in text:
            contact = text.split(':', 1)[-1].strip() if ':' in text else text
            current_supplier['contact_person'] = contact
        
        elif '主营' in text or 'Products' in text or '产品' in text:
            products = text.split(':', 1)[-1].strip() if ':' in text else text
            current_supplier['products'] = products
    
    # 最后一个供应商
    if current_supplier:
        suppliers.append(current_supplier)
    
    # 检查表格
    for table in doc.tables:
        # 表格可能包含供应商信息
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            
            # 检测是否是供应商信息行
            for cell in cells:
                if '公司名称' in cell or 'Company' in cell:
                    # 解析表格中的供应商信息
                    pass
    
    return suppliers

def main():
    print('🚀 开始提取展会供应商数据...\n')
    
    word_dir = '/Users/rabbit-y/Documents/消防装备供应商资料/供应商/消防设备供应商目录（2025）/消防设备供应商目录word版（清晰）'
    
    doc_files = [
        f'{word_dir}/消防设备供应商目录（2025）_1-100.docx',
        # 先处理第一个文件测试
    ]
    
    all_suppliers = []
    
    for doc_file in doc_files:
        print(f'📄 处理: {doc_file.split("/")[-1]}')
        suppliers = extract_suppliers_from_docx(doc_file)
        all_suppliers.extend(suppliers)
        print(f'   提取 {len(suppliers)} 家供应商\n')
    
    print(f'✅ 总共提取: {len(all_suppliers)} 家供应商\n')
    
    # 显示前 5 个
    print('📋 示例数据:\n')
    for i, supplier in enumerate(all_suppliers[:5]):
        print(f'{i+1}. {supplier.get("company_name", "(无名称)")}')
        if supplier.get('phone'):
            print(f'   电话: {supplier["phone"]}')
        if supplier.get('email'):
            print(f'   邮箱: {supplier["email"]}')
        if supplier.get('website'):
            print(f'   网址: {supplier["website"]}')
        print('')
    
    # 保存
    output_path = '/Users/rabbit-y/firegrid/backend/scripts/exhibition-suppliers-raw.json'
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(all_suppliers, f, ensure_ascii=False, indent=2)
    
    print(f'💾 已保存到: {output_path}')

if __name__ == '__main__':
    main()
